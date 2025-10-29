# utils/docx_exporter.py
"""
Export structured notes to DOCX format for download.
Compatible with both PDF analysis and saved notes formats.
"""

from __future__ import annotations
import io
import re
from typing import Dict, List, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _safe_filename(name: str) -> str:
    """Sanitize filename for safe storage."""
    name = re.sub(r"[^A-Za-z0-9 _\-.]", "", name).strip()
    return name or f"notes_{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}"


def build_docx_from_notes(structured: Dict, title: str = "UPSC Notes") -> bytes:
    """
    Build a DOCX document from structured notes (PDF analyzer output).
    
    Args:
        structured: Dictionary containing grouped notes
        title: Document title
    
    Returns:
        bytes: DOCX file as bytes
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    timestamp = structured.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    meta = doc.add_paragraph(f"Generated on: {timestamp}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_format = meta.runs[0].font
    meta_format.size = Pt(10)
    meta_format.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Summary stats
    total_items = structured.get("total_items", 0)
    categories = structured.get("categories", [])
    
    stats = doc.add_paragraph()
    stats.add_run(f"Total Cards: {total_items} | Categories: {len(categories)}")
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Content
    grouped = structured.get("grouped", {})
    
    for category, items in sorted(grouped.items(), key=lambda x: (-len(x[1]), x[0])):
        # Category heading
        doc.add_heading(category.upper().replace("_", " "), 1)
        
        for idx, item in enumerate(sorted(items, key=lambda x: x.get("relevance", 0), reverse=True), 1):
            # Card number and relevance
            card_header = doc.add_paragraph()
            card_header.add_run(f"Card #{idx} | Relevance: {item.get('relevance', 0)}/10").bold = True
            
            # Timestamp and source
            meta_info = doc.add_paragraph()
            meta_info.add_run(f"📅 {item.get('timestamp', '')} • 🔗 {item.get('source', 'N/A')}")
            meta_info_format = meta_info.runs[0].font
            meta_info_format.size = Pt(9)
            meta_info_format.color.rgb = RGBColor(100, 100, 100)
            
            # Headline (if present)
            headline = item.get("headline", "")
            if headline:
                doc.add_heading("Headline", 3)
                doc.add_paragraph(headline)
            
            # Summary
            summary = item.get("summary", "")
            if summary:
                doc.add_heading("Summary", 3)
                doc.add_paragraph(summary)
            
            # Prelims points
            prelims = item.get("prelims", []) or (item.get("deep", {}) or {}).get("prelims_points", [])
            if prelims:
                doc.add_heading("Prelims Pointers", 3)
                for p in prelims[:3]:
                    doc.add_paragraph(str(p), style='List Bullet')
            
            # Mains angles
            deep = item.get("deep", {})
            mains = deep.get("mains_angles", []) if deep else []
            if mains:
                doc.add_heading("Mains Analysis", 3)
                for m in mains[:2]:
                    doc.add_paragraph(str(m), style='List Bullet')
            
            # Interview questions
            questions = deep.get("interview_questions", []) if deep else []
            if questions:
                doc.add_heading("Interview Questions", 3)
                for q in questions[:2]:
                    doc.add_paragraph(str(q), style='List Number')
            
            # Separator
            doc.add_paragraph("_" * 80)
            doc.add_paragraph()
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"UNISOLE UPSC Notes | Generated: {timestamp}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio.read()


def export_notes_to_docx(
    notes_for_date: Dict,
    cover: bool = True,
    language: str = "en",
    cover_title: str | None = None,
) -> bytes:
    """
    Build a DOCX file from saved notes structure (API format).
    This handles the format from saved_notes.json.
    
    Args:
        notes_for_date: Dict with note data
        cover: Include cover page
        language: Language for summary (en/hi)
        cover_title: Custom cover title
    
    Returns:
        bytes: DOCX file as bytes
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Cover page
    if cover:
        title_text = cover_title or notes_for_date.get("title", "UPSC Notes")
        doc.add_heading(title_text, level=0)
        
        meta_lines = []
        date_line = notes_for_date.get("date") or notes_for_date.get("publishedAt") or ""
        if date_line:
            meta_lines.append(f"Date: {date_line}")
        
        source = notes_for_date.get("source") or notes_for_date.get("url") or ""
        if source:
            meta_lines.append(f"Source: {source}")
        
        if meta_lines:
            p = doc.add_paragraph()
            p.add_run("\n".join(meta_lines)).italic = True
        
        doc.add_page_break()

    # Helper to add a section with bullet points
    def add_section(title: str, points: List[Any]):
        if not points:
            return
            
        doc.add_heading(title, level=1)
        
        for p in points:
            if not p:
                continue
            clean = str(p).strip()
            if not clean:
                continue
            doc.add_paragraph(clean, style="List Bullet")

    # Summary (language-specific)
    if language == "hi":
        summary = notes_for_date.get("summary_hi") or ""
    else:
        summary = notes_for_date.get("summary_en") or notes_for_date.get("summary") or ""

    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    # Prelims
    add_section("Prelims Pointers", notes_for_date.get("prelims_points", []))

    # Mains angles
    add_section("Mains Angles / Essay Points", notes_for_date.get("mains_angles", []))

    # Interview questions
    add_section("Interview Questions", notes_for_date.get("interview_questions", []))

    # Schemes / Acts / Policies
    add_section("Schemes / Acts / Policies", notes_for_date.get("schemes_acts_policies", []))

    # Institutions / Dates
    add_section("Institutions", notes_for_date.get("institutions", []))
    add_section("Important Dates", notes_for_date.get("dates", []))

    # Add raw link & metadata at end
    doc.add_heading("Metadata", level=2)
    meta = []
    if notes_for_date.get("url"):
        meta.append(f"URL: {notes_for_date.get('url')}")
    if notes_for_date.get("source"):
        meta.append(f"Source: {notes_for_date.get('source')}")
    if meta:
        doc.add_paragraph("\n".join(meta))

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()